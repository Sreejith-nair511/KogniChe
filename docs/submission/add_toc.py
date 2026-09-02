"""
Inserts a Table of Contents page into the UPDATED Word document.
Inserts after the cover page (after paragraph index 19), before Section 1.
Does NOT change any existing content — only inserts the TOC page.
"""
import sys, copy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

SRC  = Path(r"c:\2026proj\Kognivera\docs\NEXORA_APS-04_Design_Submission_UPDATED.docx")
DEST = Path(r"c:\2026proj\Kognivera\docs\NEXORA_APS-04_Final_with_TOC.docx")

NAVY       = RGBColor(0x0d, 0x13, 0x21)
TEAL       = RGBColor(0x0d, 0x8a, 0x6e)
TEAL_LIGHT = RGBColor(0x4e, 0xc9, 0xa4)
MUTED      = RGBColor(0x55, 0x66, 0x77)
LIGHT_BG   = RGBColor(0xf8, 0xfa, 0xfc)
NAVY3      = RGBColor(0x1a, 0x27, 0x44)


def rgb_hex(rgb) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


# ── Complete heading catalogue (extracted from document) ───────────────────────
# Format: (level, text)   level 1 = 20pt, level 2 = 13pt
TOC_ENTRIES = [
    # ── Original 22 sections ──────────────────────────────────────────────────
    (1, "1. Executive Summary"),
    (1, "2. Problem Understanding"),
    (1, "3. Scope and Non-Scope"),
    (1, "4. User Journey"),
    (1, "5. Solution Overview"),
    (1, "6. System Architecture"),
    (1, "7. System Flow"),
    (1, "8. AI Features and Grounding"),
    (1, "9. Hybrid Retrieval"),
    (1, "10. Personalization and Reranking"),
    (1, "11. Cold Start and Session Learning"),
    (1, "12. Multilingual Intelligence"),
    (1, "13. Explainability"),
    (1, "14. Dataset and Data Usage"),
    (1, "15. Evaluation Strategy"),
    (1, "16. Business Benefits"),
    (1, "17. Technology Stack"),
    (1, "18. 24-Hour Execution Plan"),
    (1, "19. Risks and Fallbacks"),
    (1, "20. MVP Acceptance Criteria"),
    (1, "21. Final Demo Flow"),
    (1, "22. Conclusion"),
    # ── Required hackathon sections (A–N) ────────────────────────────────────
    (1, "A. Cover — Team, Problem Statement & Roles"),
    (2, "Team Members & Roles"),
    (1, "B. Problem Understanding"),
    (2, "Five Specific Failure Modes We Address"),
    (1, "C. Scope — What We Build and What We Leave Out"),
    (2, "In MVP — Built in 24 Hours"),
    (2, "Deliberately Out of Scope"),
    (1, "D. User Journey — Traveller Point of View"),
    (2, "Screen 1 — Search"),
    (2, "Screen 2 — Results & Why This"),
    (2, "Screen 3 — Interaction and Live Reranking"),
    (1, "E. Architecture Diagram"),
    (2, "Component Summary"),
    (1, "F. Flow Diagram — Search Pipeline & Feedback Loop"),
    (1, "G. Data Model Usage"),
    (2, "APS-04 Tables We Use"),
    (2, "What We Add (Additive Only — Rule R1)"),
    (1, "H. AI Features — Grounded, Measured"),
    (1, "I. Business Benefits"),
    (2, "Value to the Traveller"),
    (2, "Value to the Platform"),
    (1, "J. Technology Stack"),
    (1, "K. 24-Hour Execution Plan"),
    (1, "L. Risks and Fallbacks"),
    (1, "M. Multilingual Approach"),
    (2, "Languages Supported"),
    (2, "Technical Approach"),
    (1, "N. XR Device Declaration"),
]


def insert_paragraph_after(doc, ref_para, new_para_xml):
    """Insert new_para_xml element immediately after ref_para in the document body."""
    ref_para._p.addnext(new_para_xml)


def make_page_break_para(doc):
    """Create a paragraph XML element that is a page break."""
    p = doc.add_paragraph()
    p._p.getparent().remove(p._p)   # detach from document body
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p._p.append(r)
    # Zero spacing
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"),  "0")
    pPr.append(spacing)
    return p._p


def make_toc_title_para(doc, text):
    from docx.oxml import OxmlElement as OE
    p = doc.add_paragraph()
    p._p.getparent().remove(p._p)
    pPr = p._p.get_or_add_pPr()

    # Spacing
    spacing = OE("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"),  str(int(Pt(16).twips)))
    pPr.append(spacing)

    # Teal bottom border
    pBdr = OE("w:pBdr")
    bot = OE("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), "0D8A6E")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Run
    r = OE("w:r")
    rPr = OE("w:rPr")
    # Font
    rFonts = OE("w:rFonts")
    rFonts.set(qn("w:ascii"), "Segoe UI")
    rFonts.set(qn("w:hAnsi"), "Segoe UI")
    rPr.append(rFonts)
    # Size: 20pt = 40 half-points
    sz = OE("w:sz");    sz.set(qn("w:val"), "40");  rPr.append(sz)
    szCs = OE("w:szCs"); szCs.set(qn("w:val"), "40"); rPr.append(szCs)
    # Bold
    bold = OE("w:b"); rPr.append(bold)
    # Colour: NAVY
    color = OE("w:color"); color.set(qn("w:val"), rgb_hex(NAVY)); rPr.append(color)
    r.append(rPr)
    t = OE("w:t")
    t.text = text
    r.append(t)
    p._p.append(r)
    return p._p


def make_toc_entry_para(doc, level, text):
    """Create a TOC entry paragraph as an XML element (detached)."""
    from docx.oxml import OxmlElement as OE
    p = doc.add_paragraph()
    p._p.getparent().remove(p._p)  # detach

    pPr = p._p.get_or_add_pPr()

    # Indent for level 2
    if level == 2:
        ind = OE("w:ind")
        ind.set(qn("w:left"), str(int(Cm(0.8).twips)))
        pPr.append(ind)

    # Spacing
    spacing = OE("w:spacing")
    spacing.set(qn("w:before"), "0")
    after = str(int(Pt(3).twips)) if level == 1 else str(int(Pt(1).twips))
    spacing.set(qn("w:after"),  after)
    pPr.append(spacing)

    # Tab stop for leader dots + right-align page number at 15cm
    tabs = OE("w:tabs")
    tab = OE("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(int(Cm(14.5).twips)))
    tabs.append(tab)
    pPr.append(tabs)

    # Build run: text + tab + "–" (placeholder — Word will fill page number on open)
    def make_run(txt, bold=False, size_pt=10, color_rgb=NAVY, italic=False):
        r = OE("w:r")
        rPr = OE("w:rPr")
        rFonts = OE("w:rFonts")
        rFonts.set(qn("w:ascii"), "Segoe UI")
        rFonts.set(qn("w:hAnsi"), "Segoe UI")
        rPr.append(rFonts)
        sz_val = str(int(size_pt * 2))
        sz = OE("w:sz");    sz.set(qn("w:val"), sz_val); rPr.append(sz)
        szCs = OE("w:szCs"); szCs.set(qn("w:val"), sz_val); rPr.append(szCs)
        if bold:
            rPr.append(OE("w:b"))
        if italic:
            rPr.append(OE("w:i"))
        col = OE("w:color"); col.set(qn("w:val"), rgb_hex(color_rgb)); rPr.append(col)
        r.append(rPr)
        t = OE("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = txt
        r.append(t)
        return r

    if level == 1:
        p._p.append(make_run(text, bold=True, size_pt=10.5, color_rgb=NAVY))
    else:
        p._p.append(make_run(text, bold=False, size_pt=9.5, color_rgb=MUTED))

    # Tab character run
    r_tab = OE("w:r")
    tab_el = OE("w:tab")
    r_tab.append(tab_el)
    p._p.append(r_tab)

    # Page number placeholder run  (italic grey "–")
    p._p.append(make_run("–", bold=False, size_pt=9, color_rgb=MUTED, italic=True))

    return p._p


def make_toc_note_para(doc, text):
    from docx.oxml import OxmlElement as OE
    p = doc.add_paragraph()
    p._p.getparent().remove(p._p)
    pPr = p._p.get_or_add_pPr()
    spacing = OE("w:spacing")
    spacing.set(qn("w:before"), str(int(Pt(14).twips)))
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)
    # Teal left border
    pBdr = OE("w:pBdr")
    left_el = OE("w:left")
    left_el.set(qn("w:val"), "single")
    left_el.set(qn("w:sz"), "12")
    left_el.set(qn("w:space"), "4")
    left_el.set(qn("w:color"), "0D8A6E")
    pBdr.append(left_el)
    pPr.append(pBdr)
    # Indent
    ind = OE("w:ind"); ind.set(qn("w:left"), str(int(Cm(0.5).twips))); pPr.append(ind)
    # Background shading
    shd = OE("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"), rgb_hex(RGBColor(0xf0,0xfa,0xf6))); pPr.append(shd)
    r = OE("w:r")
    rPr = OE("w:rPr")
    rFonts = OE("w:rFonts"); rFonts.set(qn("w:ascii"),"Segoe UI"); rFonts.set(qn("w:hAnsi"),"Segoe UI"); rPr.append(rFonts)
    sz = OE("w:sz"); sz.set(qn("w:val"),"18"); rPr.append(sz)
    szCs = OE("w:szCs"); szCs.set(qn("w:val"),"18"); rPr.append(szCs)
    rPr.append(OE("w:i"))
    col = OE("w:color"); col.set(qn("w:val"), rgb_hex(RGBColor(0x33,0x41,0x55))); rPr.append(col)
    r.append(rPr)
    t = OE("w:t"); t.set(qn("xml:space"),"preserve"); t.text = text; r.append(t)
    p._p.append(r)
    return p._p


def empty_para(doc, pts_after=4):
    from docx.oxml import OxmlElement as OE
    p = doc.add_paragraph()
    p._p.getparent().remove(p._p)
    pPr = p._p.get_or_add_pPr()
    sp = OE("w:spacing")
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), str(int(Pt(pts_after).twips)))
    pPr.append(sp)
    return p._p


def insert_toc(doc):
    """
    Find the cover section end (last paragraph before '1. Executive Summary')
    and insert the entire TOC page there.
    """
    # Find paragraph index of "1. Executive Summary"
    exec_summary_para = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "1. Executive Summary":
            exec_summary_para = p
            break

    if not exec_summary_para:
        print("  Could not find '1. Executive Summary' — appending TOC at end instead")
        anchor = doc.paragraphs[-1]
        insert_after = True
    else:
        # We want to insert the TOC block immediately BEFORE the exec_summary_para
        # Strategy: insert all TOC elements before exec_summary_para._p
        anchor = exec_summary_para
        insert_after = False

    print(f"  Inserting TOC before: '{anchor.text[:60]}'")

    # Build all TOC paragraph XML elements in FORWARD order
    elements_to_insert = []

    # Page break BEFORE TOC
    elements_to_insert.append(make_page_break_para(doc))

    # TOC title
    elements_to_insert.append(make_toc_title_para(doc, "Table of Contents"))

    # Spacer after title
    elements_to_insert.append(empty_para(doc, 8))

    # TOC entries in document order
    for level, text in TOC_ENTRIES:
        elements_to_insert.append(make_toc_entry_para(doc, level, text))

    # Note at bottom of TOC
    elements_to_insert.append(make_toc_note_para(doc,
        "NOTE  Page numbers update when you open this document in Microsoft Word. "
        "Press Ctrl+A then F9 to refresh all fields."))

    # Page break AFTER TOC (so Section 1 starts on new page)
    elements_to_insert.append(make_page_break_para(doc))

    # Insert all in FORWARD order before the anchor paragraph
    # addprevious inserts before the reference element, so we insert in order
    body = anchor._p.getparent()
    ref = anchor._p

    # We collected elements in "reversed for addprevious" order above,
    # so reverse them back to forward insertion order
    elements_to_insert.reverse()
    for el in elements_to_insert:
        ref.addprevious(el)
        ref = el  # advance anchor so next element goes after current

    print(f"  Inserted {len(elements_to_insert)} elements")


def main():
    print("=" * 60)
    print("Adding Table of Contents to NEXORA Word Document")
    print("=" * 60)

    if not SRC.exists():
        print(f"ERROR: {SRC} not found")
        return

    doc = Document(str(SRC))
    print(f"Loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    insert_toc(doc)

    doc.save(str(DEST))
    size_mb = DEST.stat().st_size / (1024 * 1024)
    print(f"\nSaved: {DEST.name}  ({size_mb:.2f} MB)")
    print(f"Paragraphs now: {len(doc.paragraphs)}")
    print("=" * 60)
    print("Open in Word and press Ctrl+A then F9 to refresh page numbers.")
    print("=" * 60)


if __name__ == "__main__":
    main()
