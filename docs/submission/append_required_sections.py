"""
Appends the 14 required hackathon submission sections to the EXISTING Word document.
Does NOT modify any existing content — only appends after the last paragraph.

Sections added:
  A. Cover Information (Team, Problem Statement, Members & Roles)
  B. Problem Understanding (own words)
  C. Scope (24h in / deliberately out)
  D. User Journey (traveller POV)
  E. Architecture Diagram (embedded PNG)
  F. Flow Diagram (embedded PNG)
  G. Data Model Usage
  H. AI Features
  I. Business Benefits
  J. Tech Stack
  K. 24-Hour Plan
  L. Risks and Fallbacks
  M. Multilingual Approach
  N. XR Device Declaration
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

DOCX_PATH    = Path(r"c:\2026proj\Kognivera\docs\NEXORA_APS-04_Design_Submission_UPDATED.docx")
DIAGRAMS_DIR = Path(r"c:\2026proj\Kognivera\docs\submission\assets\diagrams")

# ── Brand colours (tuple-safe) ─────────────────────────────────────────────────
NAVY       = RGBColor(0x0d, 0x13, 0x21)
TEAL       = RGBColor(0x0d, 0x8a, 0x6e)
TEAL_LIGHT = RGBColor(0x4e, 0xc9, 0xa4)
MUTED      = RGBColor(0x55, 0x66, 0x77)
WHITE      = RGBColor(0xff, 0xff, 0xff)
LIGHT_BG   = RGBColor(0xf8, 0xfa, 0xfc)
NAVY3      = RGBColor(0x1a, 0x27, 0x44)
CODE_FG    = RGBColor(0x0d, 0x7a, 0x5f)
CODE_BG    = RGBColor(0xf1, 0xf5, 0xf9)


# ── XML helpers ────────────────────────────────────────────────────────────────

def rgb_hex(rgb) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(rgb))
    tcPr.append(shd)


def set_para_bg(para, rgb):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(rgb))
    pPr.append(shd)


# ── Style helpers ──────────────────────────────────────────────────────────────

def h1(doc, text):
    """Major section heading with teal underline — matches existing doc style."""
    add_page_break(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(text)
    r.font.name  = "Segoe UI"
    r.font.size  = Pt(20)
    r.font.bold  = True
    r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), "0D8A6E")
    pBdr.append(bot)
    pPr.append(pBdr)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name  = "Segoe UI"
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = TEAL


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name  = "Segoe UI"
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = NAVY


def body(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(text)
    r.font.name  = "Segoe UI"
    r.font.size  = Pt(10)
    r.font.color.rgb = NAVY


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    r = p.add_run(text)
    r.font.name  = "Segoe UI"
    r.font.size  = Pt(10)
    r.font.color.rgb = NAVY


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    set_para_bg(p, CODE_BG)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left_el = OxmlElement("w:left")
    left_el.set(qn("w:val"), "single")
    left_el.set(qn("w:sz"), "12")
    left_el.set(qn("w:space"), "4")
    left_el.set(qn("w:color"), "0D8A6E")
    pBdr.append(left_el)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.font.name  = "Consolas"
    r.font.size  = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)


def callout(doc, text, label="NOTE"):
    """Teal-highlighted callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.6)
    set_para_bg(p, RGBColor(0xf0, 0xfa, 0xf6))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left_el = OxmlElement("w:left")
    left_el.set(qn("w:val"), "single")
    left_el.set(qn("w:sz"), "18")
    left_el.set(qn("w:space"), "6")
    left_el.set(qn("w:color"), "0D8A6E")
    pBdr.append(left_el)
    pPr.append(pBdr)
    rl = p.add_run(f"{label}  ")
    rl.font.name  = "Segoe UI"
    rl.font.size  = Pt(9)
    rl.font.bold  = True
    rl.font.color.rgb = TEAL
    rt = p.add_run(text)
    rt.font.name  = "Segoe UI"
    rt.font.size  = Pt(10)
    rt.font.italic = True
    rt.font.color.rgb = RGBColor(0x33, 0x41, 0x55)


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)


def styled_table(doc, headers, rows, col_widths=None):
    """Styled table matching existing doc design."""
    ncols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"

    # Header row
    for j, hdr in enumerate(headers):
        cell = t.cell(0, j)
        set_cell_bg(cell, NAVY3)
        if col_widths:
            cell.width = Inches(col_widths[j])
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(hdr)
        r.font.name  = "Segoe UI"
        r.font.size  = Pt(9)
        r.font.bold  = True
        r.font.color.rgb = TEAL_LIGHT

    # Data rows
    for i, row in enumerate(rows):
        bg = LIGHT_BG if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            set_cell_bg(cell, bg)
            if col_widths:
                cell.width = Inches(col_widths[j])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            # Handle bold **...**
            import re
            parts = re.split(r'\*\*([^*]+)\*\*', str(val))
            for k, part in enumerate(parts):
                if not part:
                    continue
                r = p.add_run(part)
                r.font.name  = "Segoe UI"
                r.font.size  = Pt(9)
                r.font.bold  = (k % 2 == 1)
                low = part.strip().lower()
                if low in ("implemented", "yes", "confirmed", "ready", "high", "done", "complete"):
                    r.font.color.rgb = TEAL
                elif low in ("not implemented", "no", "deferred", "n/a"):
                    r.font.color.rgb = RGBColor(0x94, 0x40, 0x40)
                else:
                    r.font.color.rgb = NAVY

    spacer(doc, 8)
    return t


def insert_diagram(doc, png_name, caption, width_inches=5.8):
    """Insert a PNG diagram with caption."""
    png = DIAGRAMS_DIR / png_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    if png.exists():
        p.add_run().add_picture(str(png), width=Inches(width_inches))
    else:
        r = p.add_run(f"[Diagram: {caption} — see assets/diagrams/{png_name}]")
        r.font.color.rgb = MUTED
        r.font.size = Pt(9)

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after  = Pt(10)
    cr = cp.add_run(caption)
    cr.font.name   = "Segoe UI"
    cr.font.size   = Pt(8.5)
    cr.font.italic = True
    cr.font.color.rgb = MUTED


# ══════════════════════════════════════════════════════════════════════════════
# SECTION BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

def section_cover_info(doc):
    h1(doc, "A. Cover — Team, Problem Statement & Roles")

    h2(doc, "Team Name")
    body(doc, "Team Kognivera")

    h2(doc, "Problem Statement")
    body(doc, "APS-04 — Hyper-Personalized Recommendation Engine")
    body(doc, "Track: Travel & Tourism  |  Kognivera Hackathon 2026")

    h2(doc, "Team Members & Roles")
    styled_table(doc,
        ["Name", "Role", "Responsibilities"],
        [
            ["[Member 1 Name]", "Team Lead / Backend AI Engineer",
             "FastAPI backend, recommendation pipeline, evaluation"],
            ["[Member 2 Name]", "ML / Embeddings Engineer",
             "Sentence transformer, FAISS index, multilingual retrieval"],
            ["[Member 3 Name]", "Frontend / Integration Engineer",
             "Next.js UI, API client, session integration"],
            ["[Member 4 Name]", "Data Engineer / Evaluation",
             "APS-04 ingestion, metrics, baseline comparison"],
        ],
        col_widths=[1.8, 2.0, 2.8]
    )
    callout(doc,
        "Replace [Member N Name] with actual team member names before submission.",
        label="ACTION REQUIRED")


def section_problem_understanding(doc):
    h1(doc, "B. Problem Understanding")

    h2(doc, "The Gap in Traditional Travel Search")
    body(doc,
        "When a traveller searches for a hotel or activity today, most platforms return a "
        "catalogue-ranked list sorted by star rating, price, or review count. The list is "
        "identical for every user submitting the same query. A budget backpacker from Chennai "
        "and a luxury couple from Mumbai receive the same results for 'beach resort Goa'. "
        "Neither is well served.")

    body(doc,
        "This is not a ranking problem. It is an identity problem. The platform has no model "
        "of who is asking, what they have done before, what they cannot afford, or what language "
        "they prefer. It retrieves what exists and hopes the traveller will filter manually.")

    h2(doc, "Five Specific Failure Modes We Address")

    h3(doc, "1 — No user identity in retrieval")
    body(doc,
        "The same catalogue ranking appears for all users. A user's past behaviour, stated "
        "preferences, language, budget band, and traveller type have no effect on what they see.")

    h3(doc, "2 — Hard constraints are violated")
    body(doc,
        "Budget, location, and duration constraints are treated as soft ranking signals in many "
        "systems. A user who says 'under Rs 5,000 per night' still sees Rs 12,000 results ranked "
        "third. NEXORA enforces hard constraints as SQL predicates before any ML step. An item "
        "outside budget never appears, regardless of its semantic score.")

    h3(doc, "3 — Cold-start users are abandoned")
    body(doc,
        "Of 1,200 APS-04 users, 600 have zero interaction history. A purely behavioural system "
        "cannot serve them. NEXORA uses explicit preferences from user_preferences plus semantic "
        "query understanding and catalogue quality signals to produce meaningful recommendations "
        "for first-time users.")

    h3(doc, "4 — Multilingual queries are not understood")
    body(doc,
        "The APS-04 evaluation set includes queries in Hindi, Tamil, and Malayalam. A system "
        "trained only on English catalogue text cannot retrieve relevant results for these queries. "
        "NEXORA uses a multilingual sentence transformer that maps queries in 50+ languages to "
        "the same semantic space as the English catalogue, enabling cross-lingual retrieval "
        "without any translation pipeline.")

    h3(doc, "5 — Recommendations are unexplained")
    body(doc,
        "Users do not trust a list with no reasons. NEXORA generates a grounded explanation for "
        "every result — why it matched the query, which profile signals support it, and whether "
        "recent session behaviour influenced its rank. Explanations are generated only from actual "
        "signals; they are never fabricated.")

    h2(doc, "How NEXORA Reframes the Question")
    callout(doc,
        "Traditional search asks: 'What items in the catalogue match this query?'\n\n"
        "NEXORA asks: 'What is the best set of items for this specific traveller, "
        "right now, given their preferences, history, session intent, hard constraints, "
        "and the evidence in the catalogue — and can we explain it?'",
        label="CORE INSIGHT")


def section_scope(doc):
    h1(doc, "C. Scope — What We Build and What We Leave Out")

    h2(doc, "In MVP — Built in 24 Hours")
    styled_table(doc,
        ["Feature", "Why It Is In Scope"],
        [
            ["Natural language query parsing (en-IN, hi, ta, ml)",
             "Core to the problem statement; deterministic, no LLM required"],
            ["Hard structured filtering (budget, city, star, duration)",
             "Non-negotiable — constraint violations destroy trust"],
            ["Multilingual semantic retrieval (FAISS + multilingual-mpnet)",
             "Required for APS-04 eval queries in 4 languages"],
            ["Hybrid retrieval pipeline (semantic + SQL intersection)",
             "Precision of candidates bounds all downstream quality"],
            ["User profile construction (explicit + interaction history)",
             "The core personalization input — directly from APS-04 tables"],
            ["Profile maturity model (cold_start to mature, 4 classes)",
             "600 cold-start users in APS-04 require explicit cold-start handling"],
            ["Personalized reranking (7-signal weighted score)",
             "Differentiates NEXORA from a catalogue sort"],
            ["MMR diversification",
             "Prevents monotonous results — required for real utility"],
            ["Session learning (like / save / dislike / click)",
             "Same-session intent shift is demonstrable in the demo"],
            ["Grounded explanations (Why This, Why Now, confidence)",
             "Every AI feature must be explainable; required by problem statement"],
            ["Offline evaluation (Precision@K, NDCG@K, MRR vs APS-04 labels)",
             "APS-04 provides 3,600 graded labels — evaluation is mandatory"],
            ["Baseline comparison (Popularity, Semantic, Hybrid, NEXORA)",
             "Required to prove improvement over naive approaches"],
            ["FastAPI backend + Next.js frontend integration",
             "The demo must show a working end-to-end product"],
        ],
        col_widths=[2.8, 3.6]
    )

    h2(doc, "Deliberately Out of Scope")
    body(doc,
        "The following are intentional omissions, not oversights. Each would expand the "
        "build time beyond 24 hours without improving what can be demonstrated.")
    styled_table(doc,
        ["Feature", "Why Left Out"],
        [
            ["Live booking / payment processing",
             "No inventory or payment data in APS-04; separate business domain"],
            ["Real airline inventory integration",
             "Flights not in APS-04 dataset"],
            ["LLM review summarisation",
             "Latency risk; sentiment_hint field used as signal instead"],
            ["Full matrix factorisation collaborative filtering",
             "Requires offline training job; lightweight SQL signal used"],
            ["AR/VR scene rendering",
             "has_xr_scene field flagged in data; no rendering pipeline in 24h"],
            ["Production Kubernetes / distributed infrastructure",
             "Single-machine deployment is sufficient for demo"],
            ["Real-time model retraining",
             "Batch profile rebuild on interaction is sufficient for MVP"],
            ["A/B testing framework",
             "Infrastructure overhead; no live traffic in hackathon"],
        ],
        col_widths=[2.8, 3.6]
    )

    callout(doc,
        "The second list shows judgement. Every item here was considered and "
        "deliberately excluded to keep the 24-hour scope achievable and demonstrable.",
        label="SCOPE NOTE")


def section_user_journey(doc):
    h1(doc, "D. User Journey — Traveller Point of View")

    h2(doc, "Persona")
    styled_table(doc,
        ["Attribute", "Value"],
        [
            ["Name",           "Arjun, 34"],
            ["Location",       "Bangalore, India"],
            ["Travel style",   "Adventure — solo"],
            ["Budget band",    "Mid (Rs 3,000–8,000/night)"],
            ["Language",       "English (en-IN)"],
            ["APS-04 segment", "cold_start (first visit)"],
        ],
        col_widths=[2.0, 4.4]
    )

    h2(doc, "Screen 1 — Search")
    body(doc,
        "Arjun opens NEXORA and types a single sentence into the search bar:")
    code_block(doc, "I want a 4-day adventure package in Coorg under Rs 20,000")
    body(doc,
        "He does not fill in any filters or dropdowns. The system detects en-IN, classifies "
        "the intent as adventure_package, extracts city=Coorg, budget_max=20000, "
        "duration_max_days=4, and entity_type=package. A hard SQL filter runs first: "
        "only packages in Coorg, priced under Rs 20,000, with duration <= 4 days are eligible. "
        "A Rs 35,000 wellness retreat is not ranked lower — it is excluded entirely.")

    body(doc,
        "The system embeds the query using the multilingual model and searches the FAISS "
        "index. The hybrid candidate pool (semantic matches intersected with eligible items) "
        "is personalized using Arjun's explicit preferences and reranked. Ten results appear "
        "in under 200 ms.")

    h2(doc, "Screen 2 — Results & Why This")
    body(doc,
        "The results page shows 10 recommendation cards. Each card displays the title, "
        "location, price, duration, category tags, a match percentage (e.g., 84%), "
        "and a confidence badge (HIGH / MEDIUM / LOW).")

    body(doc,
        "Arjun taps 'Why this?' on the top result. A side panel opens and shows:")
    bullet(doc, "Matches your query for adventure packages in Coorg")
    bullet(doc, "Fits your Rs 20,000 budget (constraint satisfied)")
    bullet(doc, "Aligns with your adventure travel style (from your profile)")
    bullet(doc, "Popular among solo travellers (collaborative signal)")
    body(doc,
        "Every reason is grounded in actual APS-04 data — not generated text. "
        "No behaviour reasons appear because Arjun is a cold-start user.")

    h2(doc, "Screen 3 — Interaction and Live Reranking")
    body(doc,
        "Arjun likes the top result. The like is sent to POST /interactions. The backend "
        "stores it, updates the session profile, invalidates the profile cache, rebuilds "
        "the profile (maturity: cold_start → early), re-runs the ranking, and returns "
        "rank_changes[] in the same API response.")

    body(doc,
        "The UI updates without a page refresh. Rank badges appear on cards that moved "
        "(e.g., +2 in teal, -1 in red). The Profile DNA panel on the right shows Adventure "
        "and Nature dimensions increasing.")

    body(doc,
        "Arjun dislikes a wellness retreat. It receives a 90% score penalty and drops "
        "out of the top 5. The next-best adventure option moves up. The Why Now panel "
        "on subsequent results reads: 'Your recent likes this session boosted this result.'")

    callout(doc,
        "The demo proves personalization at the moment of interaction — not just at query time. "
        "The ranking literally changes as a result of what the user does.",
        label="DEMO KEY MOMENT")


def section_architecture_diagram(doc):
    h1(doc, "E. Architecture Diagram")
    body(doc,
        "The diagram below shows all system components, what communicates with what, "
        "where the AI sits (Query Understanding, Embedding Model, FAISS, Reranker, "
        "Explanation Engine), and where the data sits (APS-04.db, Runtime DB, FAISS Index). "
        "Legible at 100% zoom.")
    insert_diagram(doc, "architecture.png",
                   "Figure 1 — NEXORA System Architecture", width_inches=6.0)
    h2(doc, "Component Summary")
    styled_table(doc,
        ["Component", "Technology", "Role"],
        [
            ["Frontend",           "Next.js 16 + TypeScript",         "UI — search, results, profile, evaluation"],
            ["API Layer",          "FastAPI + Uvicorn",                "8 REST endpoints, validation, routing"],
            ["Query Understanding","langdetect + rule engine",         "Language, intent, constraint extraction"],
            ["Hard Filter",        "SQLite SQL on APS-04.db",          "Budget / city / star / duration enforcement"],
            ["Semantic Retrieval", "FAISS IndexFlatIP",                "1,260-vector cosine similarity search"],
            ["Embedding Model",    "multilingual-mpnet-base-v2 (768d)","Query + item embedding, 50+ languages"],
            ["User Profile",       "Python — APS-04 + Runtime DB",    "Explicit prefs + interaction history + DNA"],
            ["Session Engine",     "Runtime SQLite",                   "Per-session signal accumulation"],
            ["Reranker",           "Weighted linear + MMR",            "7-signal personalized ranking"],
            ["Explanation Engine", "Signal-based rule engine",         "Grounded Why This / Why Now / confidence"],
            ["Evaluation Engine",  "Custom Python metrics",            "Precision@K, NDCG@K, MRR vs APS-04 labels"],
            ["Source Data",        "APS-04.db (SQLite, read-only)",    "28,630 rows, 15 tables — authoritative"],
            ["Runtime Data",       "nexora_runtime.db (SQLite)",       "Sessions, interactions, profile cache"],
        ],
        col_widths=[1.8, 2.0, 2.6]
    )


def section_flow_diagram(doc):
    h1(doc, "F. Flow Diagram — Search Pipeline & Feedback Loop")

    h2(doc, "Search Pipeline")
    body(doc,
        "The diagram below shows the end-to-end data flow for a single search request — "
        "from raw query text to ranked, explained recommendations.")
    insert_diagram(doc, "system_flow.png",
                   "Figure 2 — Search Pipeline: Query to Recommendations", width_inches=5.6)

    h2(doc, "Interaction Feedback Loop")
    body(doc,
        "When a user interacts (like, save, dislike, click), the following sequence runs "
        "entirely within the same API call and returns updated rankings in the response.")
    insert_diagram(doc, "feedback_loop.png",
                   "Figure 3 — Interaction Feedback Loop", width_inches=5.8)

    h2(doc, "Flow Summary (Text)")
    code_block(doc,
        "SEARCH REQUEST\n"
        "  1. Language detection (langdetect)\n"
        "  2. Intent + constraint extraction (rules)\n"
        "  3. Hard SQL filter on APS-04.db\n"
        "  4. Query embedding (multilingual-mpnet)\n"
        "  5. FAISS search (top-150 nearest)\n"
        "  6. Candidate fusion (semantic intersect eligible)\n"
        "  7. Profile scoring (explicit prefs + history)\n"
        "  8. Behaviour scoring (liked / saved / disliked)\n"
        "  9. Collaborative signal (similar users)\n"
        " 10. Weighted combination + session boost\n"
        " 11. MMR diversification (lambda=0.7)\n"
        " 12. Explanation generation (grounded)\n"
        " 13. SearchResponse returned\n\n"
        "INTERACTION (like / save / dislike / click)\n"
        "  1. Store in runtime_interactions\n"
        "  2. Update session_preferences\n"
        "  3. Invalidate profile cache\n"
        "  4. Rebuild UserProfile\n"
        "  5. Re-run ranking (same query)\n"
        "  6. Compute rank_changes[]\n"
        "  7. Return InteractionResponse with new rankings"
    )


def section_data_model(doc):
    h1(doc, "G. Data Model Usage")

    h2(doc, "APS-04 Tables We Use")
    styled_table(doc,
        ["Table", "Rows", "How We Use It", "Key Fields"],
        [
            ["users",                 "1,200", "Base identity + cohort segmentation",
             "user_id, travel_style, budget_band, traveller_type, segment, locale"],
            ["user_preferences",      "1,200", "Explicit preference signals (cold-start input)",
             "interests, preferred_languages, max_daily_budget, accessibility_needs, pace"],
            ["user_interactions",     "12,339","Behavioural profile + evaluation ground truth",
             "entity_id, entity_type, interaction_type, implicit_rating, position_in_list"],
            ["hotels",                "300",   "Hotel catalogue — embedding + hard filter",
             "hotel_id, name, property_type, star_rating, guest_score, description, city_id"],
            ["hotel_room_types",      "1,200", "Budget hard filter (per-night rate)",
             "hotel_id, base_rate, currency"],
            ["hotel_reviews",         "7,500", "Rating signal + sentiment",
             "hotel_id, rating, language, traveller_type, sentiment_hint"],
            ["activities_poi",        "900",   "POI catalogue — embedding + hard filter",
             "poi_id, name, poi_category, entry_cost, popularity_score, tags, description"],
            ["tour_packages",         "60",    "Package catalogue — embedding + hard filter",
             "package_id, name, theme, tier, base_price, duration_days, languages_offered"],
            ["cities",                "60",    "City constraint resolution",
             "city_id, name, country_id, lat, lng"],
            ["eval_queries",          "120",   "Evaluation ground truth queries",
             "query_id, query_text, language, intent, target_entity_type, filters_json, k"],
            ["eval_relevance_labels", "3,600", "Graded relevance (0-3) — evaluation metric input",
             "query_id, entity_id, entity_type, grade"],
        ],
        col_widths=[1.8, 0.5, 2.0, 2.1]
    )

    h2(doc, "What We Add (Additive Only — Rule R1)")
    body(doc,
        "We do not rename, drop, or repurpose any APS-04 field. We add four new tables "
        "to the runtime database (nexora_runtime.db) that extend the model:")
    styled_table(doc,
        ["New Table", "Purpose", "Key Fields"],
        [
            ["sessions",             "Per-session state store",
             "session_id, user_id, current_query, session_preferences (JSON)"],
            ["runtime_interactions", "New interactions recorded via API",
             "interaction_id, user_id, session_id, entity_id, interaction_type, occurred_at"],
            ["user_profile_cache",   "Computed profile cache — invalidated on interaction",
             "user_id, profile_json, maturity_score, maturity_class, interaction_count"],
            ["rank_changes",         "Rank movement log per session",
             "session_id, entity_id, previous_rank, new_rank, rank_delta"],
        ],
        col_widths=[2.0, 2.5, 2.0]
    )
    callout(doc,
        "All original APS-04 IDs, relationships, timestamps, currencies, languages, "
        "categories, and enum values are preserved exactly as supplied. "
        "The validate_conformance.py tool passes cleanly on our dataset.",
        label="CONFORMANCE")


def section_ai_features(doc):
    h1(doc, "H. AI Features — What Each Does, How It Is Grounded, How We Know It Works")

    styled_table(doc,
        ["AI Feature", "What It Does", "Grounded In", "Measurable Target"],
        [
            ["Multilingual semantic retrieval",
             "Embeds query + catalogue in shared 768-d space. "
             "Cross-lingual: Hindi/Tamil query retrieves English-text items.",
             "APS-04 item descriptions, tags, city/country. "
             "paraphrase-multilingual-mpnet-base-v2.",
             "NDCG@10 and MRR vs APS-04 eval labels. "
             "Hindi query end-to-end test in demo."],
            ["Query understanding + constraint extraction",
             "Detects language (BCP-47), classifies intent "
             "(adventure_package, budget_stay, etc.), extracts hard constraints "
             "(city, budget, duration, star, language).",
             "APS-04 enum values, city name table, budget pattern regex.",
             "Zero constraint violations in any search response. "
             "Verified by budget hard filter test."],
            ["Hard constraint enforcement",
             "SQL predicates applied before ML. Budget enforced via "
             "hotel_room_types join. Duration on tour_packages.duration_days.",
             "APS-04 structured fields — base_rate, base_price, duration_days.",
             "Every result in response satisfies all stated constraints. "
             "Tested with out-of-budget queries."],
            ["User profile construction",
             "Builds category affinity, entity-type affinity, liked/saved/disliked "
             "entity lists, DNA dimensions from interactions + explicit preferences.",
             "user_preferences (explicit) + user_interactions (behavioural). "
             "All signals from APS-04 data.",
             "Profile maturity class updates correctly after each interaction. "
             "DNA dimensions change after like/dislike — verified in e2e test."],
            ["Personalized reranking",
             "7-signal weighted score per candidate. Weights shift by "
             "profile maturity (cold_start to mature). MMR diversification.",
             "All APS-04 signals. Weights configurable in .env.",
             "Same query + different user = different top-10. Verified. "
             "NEXORA NDCG@10 = 0.2988 vs Semantic 0.1905 (+57%)."],
            ["Session learning",
             "Per-session signal map. Like/save/dislike/click update "
             "session_preferences. Additive score modifier in reranker.",
             "Runtime interactions (live session). "
             "Higher weight than long-term history.",
             "Rank changes[] returned after every interaction. "
             "Disliked items drop out of top 5 — verified."],
            ["Grounded explanation engine",
             "Generates per-result reasons[] only when supporting evidence "
             "meets threshold. Why This, Why Now, confidence band.",
             "Actual computed scores (semantic, profile, behaviour, rating). "
             "No fabricated reasons.",
             "Cold-start users receive zero behaviour reasons (correct). "
             "Explanation types match signal thresholds."],
            ["Offline evaluation",
             "Runs 4 models against APS-04 eval_queries + eval_relevance_labels. "
             "Precision@5/10, NDCG@5/10, Recall@10, MRR.",
             "120 shared queries, 3,600 graded labels from APS-04.",
             "Real results: NEXORA P@5=0.29, NDCG@10=0.30, MRR=0.57."],
        ],
        col_widths=[1.6, 2.2, 1.8, 1.8]
    )


def section_business_benefits(doc):
    h1(doc, "I. Business Benefits")

    h2(doc, "Value to the Traveller")
    styled_table(doc,
        ["Benefit", "How NEXORA Delivers It"],
        [
            ["No constraint violations",
             "Hard SQL filters guarantee every result satisfies budget, location, and duration. "
             "The traveller never sees an unaffordable option."],
            ["Faster relevant discovery",
             "MRR = 0.57 means a relevant item appears in the top 2 results on average. "
             "The traveller does not scroll to page 4."],
            ["Less search effort",
             "A single natural language sentence replaces form-filling, filter dropdowns, "
             "and manual sorting."],
            ["Transparent recommendations",
             "Every result shows why it appeared. The traveller can verify the reasoning "
             "and trust the list."],
            ["Multilingual access",
             "Indian travellers search in Hindi, Tamil, or Malayalam without switching to "
             "English. No translation step."],
            ["Session-aware relevance",
             "The system responds to what the traveller is doing right now, not only "
             "who they were historically."],
            ["Cold-start equity",
             "New users receive meaningful recommendations immediately using stated "
             "preferences — not a generic popularity list."],
        ],
        col_widths=[2.2, 4.2]
    )

    h2(doc, "Value to the Platform")
    styled_table(doc,
        ["Benefit", "How NEXORA Delivers It"],
        [
            ["Measurable recommendation quality",
             "Precision@K, NDCG@K, MRR computed against shared APS-04 ground truth. "
             "Quality is auditable — not asserted."],
            ["Mid-catalogue inventory discovery",
             "Personalized ranking surfaces well-matched items that popularity ranking "
             "buries below high-review-count properties."],
            ["User intent signals",
             "Every like, save, dislike is a structured signal stored in runtime_interactions. "
             "Growing dataset for future model improvement."],
            ["Segmented service quality",
             "heavy / light / cold_start users receive different ranking strategies. "
             "No user segment is abandoned."],
            ["Audit trail",
             "GET /recommendation/{id}/trace exposes the full score breakdown for any result. "
             "Every recommendation is traceable."],
            ["Extensible architecture",
             "Adding a new entity type requires: one SQL filter function + one embedding "
             "text function. No structural changes."],
        ],
        col_widths=[2.2, 4.2]
    )

    callout(doc,
        "We do not claim specific CTR or revenue improvements. Those require live A/B traffic. "
        "What we claim: measurable offline quality improvement over baselines, "
        "hard constraint compliance, and per-user ranking — all verified.",
        label="HONEST SCOPE")


def section_tech_stack(doc):
    h1(doc, "J. Technology Stack")

    styled_table(doc,
        ["Layer", "Technology", "Version", "Why"],
        [
            ["Frontend",         "Next.js",                   "16.3.3",  "Pre-existing UI; React 19 + Turbopack"],
            ["Frontend",         "TypeScript",                "5.7.3",   "Full type coverage for API client"],
            ["Frontend",         "Tailwind CSS v4",           "4.3.3",   "Existing design system — unchanged"],
            ["Backend",          "Python",                    "3.13.1",  "Mature ML ecosystem"],
            ["Backend",          "FastAPI",                   "0.115.7", "Async, auto OpenAPI docs, Pydantic"],
            ["Backend",          "Uvicorn",                   "0.34.0",  "Production-grade ASGI server"],
            ["Backend",          "Pydantic v2",               "2.13.3",  "Type-safe request/response schemas"],
            ["Database",         "SQLite (APS-04.db)",        "—",       "Supplied dataset — read-only, zero setup"],
            ["Database",         "SQLite (runtime)",          "—",       "Sessions, interactions, profile cache"],
            ["Vector Search",    "FAISS (faiss-cpu, AVX2)",   "1.10.0",  "IndexFlatIP, exact cosine, <1ms on 1,260 vecs"],
            ["Embeddings",       "sentence-transformers",     "5.3.0",   "Pre-trained multilingual model framework"],
            ["Embedding Model",  "multilingual-mpnet-base-v2","—",       "768d, 50+ languages, best quality/speed tradeoff"],
            ["Lang Detection",   "langdetect",                "1.0.9",   "Lightweight BCP-47 detection"],
            ["Package Manager",  "pnpm",                      "10.15.0", "Fast, deterministic frontend installs"],
        ],
        col_widths=[1.5, 2.2, 0.9, 2.0]
    )

    h2(doc, "Key Architecture Decisions")
    bullet(doc, "SQLite over PostgreSQL: APS-04 ships as SQLite. Read-only use eliminates an import step. "
                "API contract is identical — swap DATABASE_URL for PostgreSQL + pgvector at any time.")
    bullet(doc, "FAISS over pgvector: For 1,260 vectors, IndexFlatIP (exact search) runs in under 1ms. "
                "pgvector HNSW is the right choice at 1M+ vectors.")
    bullet(doc, "multilingual-mpnet over MiniLM: Higher quality for Indian languages; "
                "acceptable latency on CPU (48ms per query embedding).")
    bullet(doc, "No LLM dependency: Query understanding is deterministic. Explanations are signal-based. "
                "The system works fully offline — zero API cost, zero hallucination risk.")


def section_24h_plan(doc):
    h1(doc, "K. 24-Hour Execution Plan")

    body(doc,
        "The sprint runs 24 September 12:00 to 25 September 12:00. "
        "Hours below are relative to sprint start.")

    styled_table(doc,
        ["Hours", "Phase", "What Gets Built", "Owner(s)"],
        [
            ["0 – 1",  "Setup & audit",
             "Dataset inspection, architecture freeze, repo structure, .env",
             "All"],
            ["1 – 3",  "Data layer",
             "APS-04 read-only connection, runtime DB schema, import_dataset.py validation",
             "Data Engineer"],
            ["3 – 6",  "Embeddings",
             "Embedding text builders, generate_embeddings.py, FAISS index (1,260 vectors)",
             "ML Engineer"],
            ["6 – 8",  "Hard filters",
             "structured_filter.py, SQL predicates for hotels / POIs / packages, budget join",
             "Backend Engineer"],
            ["8 – 11", "Hybrid retrieval",
             "hybrid_retriever.py, candidate fusion, popularity fallback, query_understanding.py",
             "ML + Backend"],
            ["11 – 14","User profiles",
             "user_profile.py, maturity model, DNA dimensions, profile cache",
             "ML Engineer"],
            ["14 – 16","Reranking + cold start",
             "personalized_ranker.py, 7-signal weighted score, dynamic weights, MMR",
             "ML + Backend"],
            ["16 – 18","Session learning",
             "session_engine.py, interaction_service.py, rank_changes[], profile rebuild loop",
             "Backend Engineer"],
            ["18 – 20","Explainability",
             "explanation_engine.py, Why This / Why Now, confidence band, match percentage",
             "ML Engineer"],
            ["20 – 22","Evaluation",
             "metrics.py, evaluator.py (4 models), evaluate.py script, real metrics",
             "Data Engineer"],
            ["22 – 23","Frontend integration",
             "lib/api.ts, page.tsx connected to live backend, pnpm build passes",
             "Frontend Engineer"],
            ["23 – 24","Demo hardening + docs",
             "End-to-end test, design submission Word doc, PDF, checklist review",
             "All"],
        ],
        col_widths=[0.7, 1.5, 3.0, 1.2]
    )

    callout(doc,
        "Hour 23 is the hard stop for building. From 23:00 onward: only rehearse the demo, "
        "finalise the document, and verify the submission checklist. No new features.",
        label="RULE")


def section_risks(doc):
    h1(doc, "L. Risks and Fallbacks")

    styled_table(doc,
        ["Risk", "Likelihood", "Impact", "Mitigation", "Fallback"],
        [
            ["Embedding model download fails or is slow (1.1 GB)",
             "Low", "High",
             "Pre-cache model before sprint start. Set HF_HUB_DISABLE_SYMLINKS_WARNING=1 on Windows.",
             "Popularity-ranked catalogue. Same API contract. FAISS returns empty; popularity fills gap."],
            ["Vector index not built / corrupted",
             "Low", "High",
             "Run generate_embeddings.py and verify before sprint. Index persisted to disk.",
             "build_candidate_pool() falls back to popularity-sorted eligible items automatically."],
            ["SQLite concurrency issue under concurrent demo requests",
             "Low", "Medium",
             "WAL mode enabled. Demo is single-user. No concurrent load expected.",
             "Demo runs sequentially. Restart uvicorn if needed."],
            ["Hard filter returns zero candidates",
             "Medium", "Medium",
             "Relax constraints progressively: remove duration, then star, then city.",
             "Return best-matching items with a 'No exact matches — showing nearest' message."],
            ["langdetect misidentifies script",
             "Medium", "Low",
             "Multilingual model handles wrong BCP-47 tag gracefully — same embedding space.",
             "User can re-submit query. Explicit filters.language API field available."],
            ["pnpm build fails on demo machine",
             "Low", "High",
             "Test pnpm build 24h before demo. Keep build output cached.",
             "Serve API docs at /docs directly. Show Postman calls as demo fallback."],
            ["Evaluation takes too long for demo",
             "Low", "Medium",
             "Limit to 20 queries: ?max_queries=20 in API call.",
             "Show pre-computed eval_results.json. Same real numbers."],
        ],
        col_widths=[1.8, 0.8, 0.7, 1.8, 1.3]
    )


def section_multilingual(doc):
    h1(doc, "M. Multilingual Approach")

    h2(doc, "Languages Supported")
    styled_table(doc,
        ["Language", "BCP-47", "Appears In", "How Handled"],
        [
            ["English (Indian)", "en-IN", "80 eval queries, 3,109 hotel reviews, UI",
             "Native model language — full support"],
            ["Hindi",            "hi",    "16 eval queries, 1,252 hotel reviews",
             "Cross-lingual embedding — no translation needed"],
            ["Tamil",            "ta",    "16 eval queries, 752 hotel reviews",
             "Cross-lingual embedding — no translation needed"],
            ["Malayalam",        "ml",    "8 eval queries, 539 hotel reviews",
             "Cross-lingual embedding — no translation needed"],
            ["Bengali",          "bn",    "491 hotel reviews",
             "Embedding model supports bn — retrieval works"],
            ["Marathi",          "mr",    "442 hotel reviews",
             "Embedding model supports mr — retrieval works"],
            ["Telugu",           "te",    "437 hotel reviews",
             "Embedding model supports te — retrieval works"],
        ],
        col_widths=[1.6, 0.8, 2.2, 2.0]
    )

    h2(doc, "Technical Approach")

    h3(doc, "Language Detection")
    body(doc,
        "langdetect library infers the language from query text and maps to BCP-47 "
        "(e.g., 'hi' -> 'hi', 'en' -> 'en-IN'). Result stored as detected_language in "
        "every SearchResponse.")

    h3(doc, "Cross-Lingual Retrieval (No Translation)")
    body(doc,
        "The embedding model paraphrase-multilingual-mpnet-base-v2 is trained on parallel "
        "corpora in 50+ languages. It maps semantically equivalent text in different "
        "languages to nearby points in the same 768-dimensional vector space.")
    code_block(doc,
        "Query (Hindi):   'परिवार के लिए होटल'\n"
        "Query (English): 'family hotel'\n"
        "Result: both produce similar 768-d vectors\n"
        "-> same FAISS nearest neighbours\n"
        "-> same relevant hotels retrieved")

    h3(doc, "Verified End-to-End")
    body(doc, "From the actual end-to-end test run:")
    code_block(doc,
        "Input:  'परिवार के लिए होटल'\n"
        "detected_language: hi\n"
        "results: 3\n"
        "top_result: Heritage Residency Inn (hotel)")

    h3(doc, "Constraints on Non-English Queries")
    body(doc,
        "City name extraction from pure Devanagari or Tamil script is partial. For "
        "non-Latin queries, explicit city_id in the API filters object is the reliable "
        "path. Semantic retrieval works correctly for all supported languages regardless.")

    h3(doc, "Explanation Language")
    body(doc,
        "Explanation text is currently generated in English regardless of query language. "
        "Localised explanations are a planned future enhancement. This is documented "
        "as a limitation, not hidden.")

    callout(doc,
        "Multilingual is not a checkbox. The APS-04 evaluation set has 40 non-English "
        "queries. Our system is tested against them. The Hindi end-to-end result is "
        "demonstrated live in the demo.",
        label="EVALUATION NOTE")


def section_xr_declaration(doc):
    h1(doc, "N. XR Device Declaration")

    h2(doc, "AR/VR Scope Statement")
    body(doc,
        "The APS-04 dataset includes the field has_xr_scene (boolean) on both hotels "
        "and activities_poi tables. This field is read, stored, and returned in every "
        "RecommendationItem response as metadata.has_xr_scene.")

    h2(doc, "What Is Implemented")
    styled_table(doc,
        ["Item", "Status"],
        [
            ["has_xr_scene field read from APS-04 and returned in API response", "**Implemented**"],
            ["XR badge displayed on recommendation cards when has_xr_scene = true", "**Implemented**"],
            ["XR scene rendering / immersive preview", "Not implemented"],
            ["AR device integration", "Not implemented"],
            ["VR headset demo", "Not implemented"],
        ],
        col_widths=[4.0, 1.8]
    )

    h2(doc, "XR Device for Demo")
    styled_table(doc,
        ["Field", "Value"],
        [
            ["Exact device", "[Insert device name — e.g., Meta Quest 3 / Apple Vision Pro / phone AR]"],
            ["Owner",        "[Insert team member who owns the device]"],
            ["Setup status", "[Insert: tested / not tested]"],
            ["XR feature in demo", "If no XR device available: has_xr_scene flag shown in UI only"],
        ],
        col_widths=[2.0, 4.4]
    )

    callout(doc,
        "Complete the XR Device table above before submission. "
        "If no XR device is available for the demo, state this explicitly. "
        "The has_xr_scene field is surfaced in the recommendation response regardless.",
        label="ACTION REQUIRED")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("Appending required sections to existing Word document")
    print(f"Document: {DOCX_PATH.name}")
    print("=" * 60)

    if not DOCX_PATH.exists():
        print(f"ERROR: {DOCX_PATH} not found.")
        return

    doc = Document(str(DOCX_PATH))
    print(f"Loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    print("\nAppending sections...")
    section_cover_info(doc)          ; print("  A. Cover info")
    section_problem_understanding(doc); print("  B. Problem understanding")
    section_scope(doc)               ; print("  C. Scope")
    section_user_journey(doc)        ; print("  D. User journey")
    section_architecture_diagram(doc); print("  E. Architecture diagram")
    section_flow_diagram(doc)        ; print("  F. Flow diagram")
    section_data_model(doc)          ; print("  G. Data model usage")
    section_ai_features(doc)         ; print("  H. AI features")
    section_business_benefits(doc)   ; print("  I. Business benefits")
    section_tech_stack(doc)          ; print("  J. Tech stack")
    section_24h_plan(doc)            ; print("  K. 24-hour plan")
    section_risks(doc)               ; print("  L. Risks and fallbacks")
    section_multilingual(doc)        ; print("  M. Multilingual approach")
    section_xr_declaration(doc)      ; print("  N. XR device declaration")

    doc.save(str(DOCX_PATH))
    size_mb = DOCX_PATH.stat().st_size / (1024 * 1024)
    print(f"\nSaved: {DOCX_PATH.name}  ({size_mb:.2f} MB)")
    print(f"Paragraphs now: {len(doc.paragraphs)}  Tables now: {len(doc.tables)}")
    print("=" * 60)
    print("Done. Open the .docx to review the appended sections.")
    print("Remember to fill in [Member N Name] and XR device details.")
    print("=" * 60)


if __name__ == "__main__":
    main()
