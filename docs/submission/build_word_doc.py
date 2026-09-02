"""
NEXORA Design Submission — Word Document Builder
Produces a professionally designed .docx with:
  - Branded cover page (navy/teal)
  - Styled headings (3 levels)
  - Real tables with alternating rows
  - Code blocks
  - Embedded diagram images
  - Page headers/footers
  - Section numbering
  - All 22 source documents compiled in order
"""
import sys, re
from pathlib import Path
from copy import deepcopy

# Force UTF-8
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants  # noqa

# ── Brand colours ──────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x0d, 0x13, 0x21)
NAVY2      = RGBColor(0x13, 0x1c, 0x2e)
NAVY3      = RGBColor(0x1a, 0x27, 0x44)
TEAL       = RGBColor(0x0d, 0x8a, 0x6e)
TEAL_LIGHT = RGBColor(0x4e, 0xc9, 0xa4)
WHITE      = RGBColor(0xff, 0xff, 0xff)
OFFWHITE   = RGBColor(0xf0, 0xf4, 0xf8)
MUTED      = RGBColor(0x55, 0x66, 0x77)
LIGHT_BG   = RGBColor(0xf8, 0xfa, 0xfc)
BORDER     = RGBColor(0xcc, 0xd6, 0xe0)
CODE_BG    = RGBColor(0xf1, 0xf5, 0xf9)
CODE_FG    = RGBColor(0x0d, 0x7a, 0x5f)
TABLE_HEAD = RGBColor(0xee, 0xf2, 0xf7)
TABLE_ALT  = RGBColor(0xfa, 0xfd, 0xff)
RED_SOFT   = RGBColor(0xfe, 0xf2, 0xf2)

SUBMISSION_DIR = Path(__file__).parent
OUTPUT_DOCX = SUBMISSION_DIR.parent / "NEXORA_APS-04_Design_Submission.docx"
DIAGRAMS_DIR = SUBMISSION_DIR / "assets" / "diagrams"

DOCUMENTS = [
    "00_COVER.md", "01_EXECUTIVE_SUMMARY.md", "02_PROBLEM_UNDERSTANDING.md",
    "03_SCOPE_AND_NON_SCOPE.md", "04_USER_JOURNEY.md", "05_SOLUTION_OVERVIEW.md",
    "06_ARCHITECTURE.md", "07_SYSTEM_FLOW.md", "08_AI_FEATURES_AND_GROUNDING.md",
    "09_HYBRID_RETRIEVAL.md", "10_PERSONALIZATION_AND_RERANKING.md",
    "11_COLD_START_AND_SESSION_LEARNING.md", "12_MULTILINGUAL_INTELLIGENCE.md",
    "13_EXPLAINABILITY.md", "14_DATASET_AND_DATA_USAGE.md",
    "15_EVALUATION_STRATEGY.md", "16_BUSINESS_BENEFITS.md",
    "17_TECH_STACK.md", "18_24_HOUR_EXECUTION_PLAN.md",
    "19_RISKS_AND_FALLBACKS.md", "20_MVP_ACCEPTANCE_CRITERIA.md",
    "21_FINAL_DEMO_FLOW.md", "22_CONCLUSION.md",
]


# ── XML helpers ────────────────────────────────────────────────────────────────

def rgb_hex(rgb) -> str:
    """Convert RGBColor (tuple) to hex string like '0D1321'."""
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(rgb))
    tcPr.append(shd)


def set_para_bg(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(rgb))
    pPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), val)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)
    remove_para_spacing(p)


def remove_para_spacing(para):
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_horizontal_rule(doc, color="0D8A6E"):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    remove_para_spacing(para)
    return para


def set_doc_margins(section, top=2.0, bottom=2.0, left=2.5, right=2.5):
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


# ── Cover page ─────────────────────────────────────────────────────────────────

def build_cover(doc):
    section = doc.sections[0]
    set_doc_margins(section)

    # Large navy background block via table trick
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, NAVY)
    # Remove table borders
    for side in ["top", "bottom", "left", "right"]:
        set_cell_border(cell, **{side: "FFFFFF"})

    cell.width = Inches(6.5)

    # NEXORA in huge teal
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("NEXORA")
    run.font.name = "Segoe UI"
    run.font.size = Pt(54)
    run.font.bold = True
    run.font.color.rgb = TEAL_LIGHT

    # Tagline
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("Adaptive Recommendation Intelligence")
    r2.font.name = "Segoe UI"
    r2.font.size = Pt(18)
    r2.font.color.rgb = OFFWHITE
    r2.font.bold = False

    # Subtitle line
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(28)
    r3 = p3.add_run("Hyper-Personalized Recommendation Engine")
    r3.font.name = "Segoe UI"
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0x88, 0x99, 0xaa)

    # Divider line in teal
    div = cell.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_before = Pt(0)
    div.paragraph_format.space_after = Pt(24)
    dr = div.add_run("─" * 42)
    dr.font.color.rgb = TEAL
    dr.font.size = Pt(9)

    # Positioning statement
    ps = cell.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.paragraph_format.space_before = Pt(0)
    ps.paragraph_format.space_after = Pt(4)
    r_ps1 = ps.add_run("Search tells you what exists.")
    r_ps1.font.name = "Segoe UI"
    r_ps1.font.size = Pt(13)
    r_ps1.font.italic = True
    r_ps1.font.color.rgb = RGBColor(0xb0, 0xc4, 0xd6)

    ps2 = cell.add_paragraph()
    ps2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps2.paragraph_format.space_before = Pt(0)
    ps2.paragraph_format.space_after = Pt(32)
    r_ps2 = ps2.add_run("NEXORA learns what belongs to you.")
    r_ps2.font.name = "Segoe UI"
    r_ps2.font.size = Pt(13)
    r_ps2.font.italic = True
    r_ps2.font.color.rgb = TEAL_LIGHT

    # Event details
    for line in [
        "Problem Statement: APS-04",
        "Track: Travel & Tourism",
        "Event: Kognivera Hackathon 2026",
        "Dataset: APS-04 · data model v1.1.0-rc1 · 28,630 rows · 15 tables",
    ]:
        pev = cell.add_paragraph()
        pev.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pev.paragraph_format.space_before = Pt(0)
        pev.paragraph_format.space_after = Pt(3)
        rev = pev.add_run(line)
        rev.font.name = "Segoe UI"
        rev.font.size = Pt(10)
        rev.font.color.rgb = RGBColor(0x88, 0x99, 0xaa)

    # Bottom padding
    pad = cell.add_paragraph()
    pad.paragraph_format.space_before = Pt(40)
    pad.paragraph_format.space_after = Pt(0)
    pad.add_run("")

    # Status mini-table
    doc.add_paragraph()
    status_items = [
        ("Hybrid Retrieval (semantic + structured)",  "Implemented"),
        ("User Profile Engine (explicit + implicit)", "Implemented"),
        ("Session Learning",                          "Implemented"),
        ("Personalized Reranking (7-signal, MMR)",    "Implemented"),
        ("Cold-Start Handling (600 test users)",      "Implemented"),
        ("Multilingual Queries (en-IN, hi, ta, ml)",  "Implemented"),
        ("Grounded Explanations (Why This / Now)",    "Implemented"),
        ("Offline Evaluation (NDCG, P@K, MRR)",       "Implemented"),
        ("Next.js Frontend — live backend connection","Implemented"),
    ]
    t = doc.add_table(rows=1 + len(status_items), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"

    # Header row
    hc0, hc1 = t.rows[0].cells
    for hc, txt in [(hc0, "Component"), (hc1, "Status")]:
        set_cell_bg(hc, NAVY3)
        hc.width = Inches(4.2) if hc == hc0 else Inches(1.8)
        hp = hc.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hr = hp.add_run(txt)
        hr.font.name = "Segoe UI"
        hr.font.size = Pt(9)
        hr.font.bold = True
        hr.font.color.rgb = TEAL_LIGHT
        hc.paragraphs[0].paragraph_format.space_before = Pt(4)
        hc.paragraphs[0].paragraph_format.space_after = Pt(4)

    for i, (comp, stat) in enumerate(status_items):
        row = t.rows[i + 1]
        c0, c1 = row.cells
        bg = LIGHT_BG if i % 2 == 0 else WHITE
        set_cell_bg(c0, bg)
        set_cell_bg(c1, bg)
        for c, txt, bold, color in [
            (c0, comp, False, NAVY),
            (c1, stat, True,  TEAL),
        ]:
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(txt)
            r.font.name = "Segoe UI"
            r.font.size = Pt(9)
            r.font.bold = bold
            r.font.color.rgb = color

    add_page_break(doc)


# ── Header / Footer ────────────────────────────────────────────────────────────

def add_header_footer(doc):
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("NEXORA  ·  APS-04  ·  Kognivera Hackathon 2026")
        r.font.name = "Segoe UI"
        r.font.size = Pt(8)
        r.font.color.rgb = MUTED
        # Teal bottom border on header
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "0D8A6E")
        pBdr.append(bot)
        pPr.append(pBdr)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run("NEXORA  \u00b7  APS-04  \u00b7  Kognivera Hackathon 2026     Page ")
        r1.font.name = "Segoe UI"
        r1.font.size = Pt(8)
        r1.font.color.rgb = MUTED
        # Page number as field
        r2 = fp.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        r2._r.append(fldChar1)
        r2._r.append(instrText)
        r2._r.append(fldChar2)
        r2.font.name = "Segoe UI"
        r2.font.size = Pt(8)
        r2.font.color.rgb = MUTED


# ── Styled paragraph helpers ───────────────────────────────────────────────────

def add_h1(doc, text):
    """Section heading with teal underline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = NAVY
    # Teal bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), "0D8A6E")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = TEAL
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = NAVY
    return p


def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.font.name = "Segoe UI"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = MUTED
    return p


def add_body(doc, text, italic=False):
    if not text.strip():
        return None
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    # Handle inline code `...`
    parts = re.split(r'`([^`]+)`', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
        r.font.italic = italic
        if i % 2 == 1:  # inside backticks
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = CODE_FG
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    parts = re.split(r'`([^`]+)`', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
        if i % 2 == 1:
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = CODE_FG
    return p


def add_code_block(doc, text):
    """Shaded monospace code block."""
    # Remove leading/trailing blank lines
    lines = text.strip("\n").split("\n")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    set_para_bg(p, CODE_BG)
    # Teal left border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left_el = OxmlElement("w:left")
    left_el.set(qn("w:val"), "single")
    left_el.set(qn("w:sz"), "12")
    left_el.set(qn("w:space"), "4")
    left_el.set(qn("w:color"), "0D8A6E")
    pBdr.append(left_el)
    pPr.append(pBdr)
    r = p.add_run("\n".join(lines))
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    return p


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.8)
    set_para_bg(p, RGBColor(0xf0, 0xfa, 0xf6))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left_el = OxmlElement("w:left")
    left_el.set(qn("w:val"), "single")
    left_el.set(qn("w:sz"), "18")
    left_el.set(qn("w:space"), "6")
    left_el.set(qn("w:color"), "0D8A6E")
    pBdr.append(left_el)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p


# ── Table builder ──────────────────────────────────────────────────────────────

def add_markdown_table(doc, header_row, data_rows):
    """Build a styled table from parsed markdown table rows."""
    ncols = len(header_row)
    nrows = len(data_rows) + 1
    t = doc.add_table(rows=nrows, cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    t.allow_autofit = True

    # Header
    for j, cell_text in enumerate(header_row):
        cell = t.cell(0, j)
        set_cell_bg(cell, NAVY3)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(cell_text.strip())
        r.font.name = "Segoe UI"
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = TEAL_LIGHT

    # Data rows
    for i, row in enumerate(data_rows):
        bg = TABLE_ALT if i % 2 == 0 else WHITE
        # Detect NEXORA highlight row
        is_nexora = any("nexora" in str(c).lower() or "**nexora**" in str(c).lower() for c in row)
        if is_nexora:
            bg = RGBColor(0xf0, 0xfa, 0xf6)
        for j, cell_text in enumerate(row):
            cell = t.cell(i + 1, j)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            txt = str(cell_text).strip()
            # Bold markdown **...**
            bold_parts = re.split(r'\*\*([^*]+)\*\*', txt)
            for k, part in enumerate(bold_parts):
                if not part:
                    continue
                r = p.add_run(part)
                r.font.name = "Segoe UI"
                r.font.size = Pt(9)
                r.font.bold = (k % 2 == 1)
                # colour certain status values
                low = part.lower().strip()
                if low in ("implemented", "yes", "ok", "pass", "high", "done"):
                    r.font.color.rgb = TEAL
                elif low in ("not implemented", "no", "fail", "low", "medium", "deferred"):
                    r.font.color.rgb = RGBColor(0x94, 0x40, 0x40)
                elif is_nexora and k % 2 == 1:
                    r.font.color.rgb = TEAL
                else:
                    r.font.color.rgb = NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


# ── Diagram injector ───────────────────────────────────────────────────────────

def try_insert_diagram(doc, label):
    """Insert a PNG diagram if available."""
    label_lower = label.lower()
    if "architecture" in label_lower:
        png = DIAGRAMS_DIR / "architecture.png"
        caption = "Figure 1 — System Architecture"
    elif "search pipeline" in label_lower or "system flow" in label_lower:
        png = DIAGRAMS_DIR / "system_flow.png"
        caption = "Figure 2 — Search Pipeline Flow"
    elif "feedback" in label_lower:
        png = DIAGRAMS_DIR / "feedback_loop.png"
        caption = "Figure 3 — Interaction Feedback Loop"
    else:
        return

    if not png.exists():
        add_body(doc, f"[Diagram: {caption} — see assets/diagrams/]")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    try:
        run.add_picture(str(png), width=Inches(5.8))
    except Exception:
        p.add_run(f"[{caption}]").font.color.rgb = MUTED

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(10)
    cr = cp.add_run(caption)
    cr.font.name = "Segoe UI"
    cr.font.size = Pt(8.5)
    cr.font.italic = True
    cr.font.color.rgb = MUTED


# ── Markdown parser ────────────────────────────────────────────────────────────

def parse_and_render(doc, md_text, is_cover=False):
    """Parse markdown line-by-line and render to Word elements."""
    lines = md_text.split("\n")
    i = 0
    in_code = False
    code_lines = []
    in_table = False
    table_header = []
    table_rows = []

    def flush_table():
        nonlocal in_table, table_header, table_rows
        if table_header:
            add_markdown_table(doc, table_header, table_rows)
        in_table = False
        table_header = []
        table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith("```"):
            if in_code:
                if code_lines:
                    add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Image reference ![alt](src)
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            if in_table:
                flush_table()
            alt = img_match.group(1)
            try_insert_diagram(doc, alt)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', line.strip()):
            if in_table:
                flush_table()
            add_horizontal_rule(doc)
            i += 1
            continue

        # Heading detection
        h_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if h_match:
            if in_table:
                flush_table()
            level = len(h_match.group(1))
            text  = h_match.group(2).strip()
            # Strip markdown bold/italic from headings
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*',   r'\1', text)
            if level == 1:
                if not is_cover:
                    add_page_break(doc)
                add_h1(doc, text)
            elif level == 2:
                add_h2(doc, text)
            elif level == 3:
                add_h3(doc, text)
            else:
                add_h4(doc, text)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            if in_table:
                flush_table()
            text = line[2:].strip()
            # Collect multi-line blockquotes
            bq_lines = [text]
            while i + 1 < len(lines) and lines[i + 1].startswith("> "):
                i += 1
                bq_lines.append(lines[i][2:].strip())
            add_blockquote(doc, " ".join(bq_lines))
            i += 1
            continue

        # Table row
        if line.strip().startswith("|") and "|" in line[1:]:
            # Separator row — skip
            if re.match(r'^[\|\s\-:]+$', line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if not in_table:
                in_table = True
                table_header = cells
                table_rows = []
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # Bullet / list
        bullet_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text   = bullet_match.group(2).strip()
            # Clean markdown bold
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            add_bullet(doc, text, level=indent)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^\s*\d+\.\s+(.*)', line)
        if num_match:
            add_bullet(doc, num_match.group(1).strip())
            i += 1
            continue

        # Checkbox list
        check_match = re.match(r'^\s*-\s+\[([ xX✅⚠️🔲])\]\s+(.*)', line)
        if check_match:
            mark = check_match.group(1)
            text = check_match.group(2).strip()
            symbol = "✓ " if mark.lower() in ("x", "✅") else ("⚠ " if "⚠" in mark else "○ ")
            col = TEAL if symbol == "✓ " else (RGBColor(0xd9, 0x77, 0x06) if symbol == "⚠ " else MUTED)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)
            rs = p.add_run(symbol)
            rs.font.name = "Segoe UI"
            rs.font.size = Pt(10)
            rs.font.bold = True
            rs.font.color.rgb = col
            rt = p.add_run(text)
            rt.font.name = "Segoe UI"
            rt.font.size = Pt(10)
            rt.font.color.rgb = NAVY
            i += 1
            continue

        # Empty line
        if not line.strip():
            if in_table:
                flush_table()
            i += 1
            continue

        # Front-matter / divider lines
        if line.strip() in ("---", "===") or line.startswith("---"):
            if in_table:
                flush_table()
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        if text:
            # Clean bold/italic markdown from body text
            is_italic = text.startswith("*") and text.endswith("*") and not text.startswith("**")
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            add_body(doc, text, italic=is_italic)

        i += 1

    if in_table:
        flush_table()
    if in_code and code_lines:
        add_code_block(doc, "\n".join(code_lines))


# ── Main builder ───────────────────────────────────────────────────────────────

def build():
    print("=" * 60)
    print("NEXORA Word Document Builder")
    print("=" * 60)

    doc = Document()

    # Set default font
    from docx.oxml.ns import nsmap
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10)
    style.font.color.rgb = NAVY

    # Default section margins
    section = doc.sections[0]
    set_doc_margins(section)

    print("\n[1/3] Building cover page...")
    build_cover(doc)

    print("[2/3] Compiling documents...")
    for fname in DOCUMENTS:
        path = SUBMISSION_DIR / fname
        if not path.exists():
            print(f"  SKIP: {fname}")
            continue
        content = path.read_text(encoding="utf-8").strip()
        is_cover = fname == "00_COVER.md"
        parse_and_render(doc, content, is_cover=is_cover)
        print(f"  OK: {fname}")

    print("[3/3] Adding headers and footers...")
    add_header_footer(doc)

    doc.save(str(OUTPUT_DOCX))
    size_mb = OUTPUT_DOCX.stat().st_size / (1024 * 1024)
    print(f"\nWord document: {OUTPUT_DOCX.name}")
    print(f"Size: {size_mb:.2f} MB")
    print("=" * 60)
    print("Done. Open in Microsoft Word or Google Docs to review and edit.")
    print("=" * 60)


if __name__ == "__main__":
    build()
